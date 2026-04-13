from __future__ import annotations

import csv
import io
from dataclasses import dataclass
from typing import Any

import openpyxl
import pdfplumber
from docx import Document as DocxDocument

from app.core.config import Settings, get_settings
from app.models import Document
from app.models.enums import DocumentFileType


@dataclass(slots=True, frozen=True)
class ExtractedTextChunk:
    chunk_index: int
    content: str
    source_type: str
    source_locator: dict[str, Any]
    metadata: dict[str, Any] | None = None


def _normalize_text(value: str) -> str:
    return "\n".join(line.rstrip() for line in value.splitlines()).strip()


def _window_text(
    text: str,
    *,
    chunk_size: int,
    overlap: int,
) -> list[tuple[str, dict[str, int]]]:
    normalized = _normalize_text(text)
    if not normalized:
        return []

    windows: list[tuple[str, dict[str, int]]] = []
    start = 0
    while start < len(normalized):
        end = min(len(normalized), start + chunk_size)
        chunk = normalized[start:end].strip()
        if chunk:
            windows.append((chunk, {"char_start": start, "char_end": end}))
        if end >= len(normalized):
            break
        start = max(end - overlap, start + 1)
    return windows


def _pdf_chunks(
    file_bytes: bytes,
    settings: Settings,
) -> list[ExtractedTextChunk]:
    chunks: list[ExtractedTextChunk] = []
    chunk_index = 0
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            text = _normalize_text(page.extract_text() or "")
            if not text:
                continue
            for window_text, bounds in _window_text(
                text,
                chunk_size=settings.rag_chunk_size_chars,
                overlap=settings.rag_chunk_overlap_chars,
            ):
                chunks.append(
                    ExtractedTextChunk(
                        chunk_index=chunk_index,
                        content=window_text,
                        source_type="pdf_page",
                        source_locator={"page": page_number, **bounds},
                    )
                )
                chunk_index += 1
    if not chunks:
        raise ValueError("PDF does not contain extractable text for indexing")
    return chunks


def _docx_blocks(file_bytes: bytes) -> list[str]:
    doc = DocxDocument(io.BytesIO(file_bytes))
    blocks: list[str] = []
    for paragraph_index, paragraph in enumerate(doc.paragraphs, start=1):
        text = _normalize_text(paragraph.text)
        if not text:
            continue
        blocks.append(f"[paragraph {paragraph_index}]\n{text}")

    for table_index, table in enumerate(doc.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            values = [_normalize_text(cell.text) for cell in row.cells]
            cleaned_values = [value for value in values if value]
            if cleaned_values:
                rows.append(" | ".join(cleaned_values))
        if rows:
            blocks.append(f"[table {table_index}]\n" + "\n".join(rows))
    return blocks


def _docx_chunks(
    file_bytes: bytes,
    settings: Settings,
) -> list[ExtractedTextChunk]:
    blocks = _docx_blocks(file_bytes)
    if not blocks:
        raise ValueError("DOCX does not contain extractable text for indexing")

    chunks: list[ExtractedTextChunk] = []
    chunk_index = 0
    for block_index, block in enumerate(blocks, start=1):
        for window_text, bounds in _window_text(
            block,
            chunk_size=settings.rag_chunk_size_chars,
            overlap=settings.rag_chunk_overlap_chars,
        ):
            chunks.append(
                ExtractedTextChunk(
                    chunk_index=chunk_index,
                    content=window_text,
                    source_type="docx_block",
                    source_locator={"block_index": block_index, **bounds},
                )
            )
            chunk_index += 1
    return chunks


def _serialize_row(headers: list[str], values: list[str]) -> str | None:
    pairs: list[str] = []
    for index, raw_value in enumerate(values):
        value = _normalize_text(raw_value)
        if not value:
            continue
        header = (
            headers[index]
            if index < len(headers) and headers[index]
            else f"coluna_{index + 1}"
        )
        pairs.append(f"{header}: {value}")
    if not pairs:
        return None
    return " | ".join(pairs)


def _tabular_chunks(
    *,
    sections: list[tuple[str, list[str], list[tuple[int, list[str]]]]],
    source_type: str,
    rows_per_chunk: int,
) -> list[ExtractedTextChunk]:
    chunks: list[ExtractedTextChunk] = []
    chunk_index = 0
    for section_name, headers, rows in sections:
        if not rows:
            continue
        for start in range(0, len(rows), rows_per_chunk):
            batch = rows[start : start + rows_per_chunk]
            serialized_rows: list[str] = []
            row_numbers: list[int] = []
            for row_number, values in batch:
                serialized = _serialize_row(headers, values)
                if not serialized:
                    continue
                serialized_rows.append(f"Linha {row_number}: {serialized}")
                row_numbers.append(row_number)
            if not serialized_rows:
                continue
            chunks.append(
                ExtractedTextChunk(
                    chunk_index=chunk_index,
                    content=f"{section_name}\n" + "\n".join(serialized_rows),
                    source_type=source_type,
                    source_locator={
                        "section": section_name,
                        "row_start": min(row_numbers),
                        "row_end": max(row_numbers),
                    },
                )
            )
            chunk_index += 1
    if not chunks:
        raise ValueError("Tabular document does not contain extractable rows")
    return chunks


def _xlsx_chunks(
    file_bytes: bytes,
    settings: Settings,
) -> list[ExtractedTextChunk]:
    workbook = openpyxl.load_workbook(
        io.BytesIO(file_bytes), data_only=True, read_only=True
    )
    sections: list[tuple[str, list[str], list[tuple[int, list[str]]]]] = []
    for sheet in workbook.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue
        header_values = [_normalize_text(str(value or "")) for value in rows[0]]
        data_rows: list[tuple[int, list[str]]] = []
        for row_offset, row in enumerate(rows[1:], start=2):
            values = ["" if value is None else str(value) for value in row]
            if any(_normalize_text(value) for value in values):
                data_rows.append((row_offset, values))
        sections.append((f"Sheet {sheet.title}", header_values, data_rows))
    return _tabular_chunks(
        sections=sections,
        source_type="xlsx_rows",
        rows_per_chunk=settings.rag_tabular_rows_per_chunk,
    )


def _csv_chunks(
    file_bytes: bytes,
    settings: Settings,
) -> list[ExtractedTextChunk]:
    decoded = file_bytes.decode("utf-8-sig", errors="ignore")
    reader = list(csv.reader(io.StringIO(decoded)))
    if not reader:
        raise ValueError("CSV does not contain extractable rows")
    headers = [_normalize_text(value) for value in reader[0]]
    rows: list[tuple[int, list[str]]] = []
    for row_index, row in enumerate(reader[1:], start=2):
        if any(_normalize_text(value) for value in row):
            rows.append((row_index, row))
    return _tabular_chunks(
        sections=[("CSV", headers, rows)],
        source_type="csv_rows",
        rows_per_chunk=settings.rag_tabular_rows_per_chunk,
    )


def extract_document_text_chunks(
    document: Document,
    file_bytes: bytes,
    settings: Settings | None = None,
) -> list[ExtractedTextChunk]:
    current_settings = settings or get_settings()
    if document.file_type == DocumentFileType.PDF:
        return _pdf_chunks(file_bytes, current_settings)
    if document.file_type == DocumentFileType.DOCX:
        return _docx_chunks(file_bytes, current_settings)
    if document.file_type == DocumentFileType.XLSX:
        return _xlsx_chunks(file_bytes, current_settings)
    if document.file_type == DocumentFileType.CSV:
        return _csv_chunks(file_bytes, current_settings)
    raise ValueError(f"Unsupported file type for indexing: {document.file_type.value}")
