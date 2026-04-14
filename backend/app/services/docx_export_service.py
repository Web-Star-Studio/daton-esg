"""Render a ``Report`` (plus its project) as an editable .docx file.

Uses python-docx (already in deps). Scope kept intentionally minimal:
  - cover paragraph with org name and year
  - H1 per major section, H2 for optional subsections emitted by the LLM as
    markdown ``## ...`` lines
  - paragraphs, bullet lists, simple markdown tables, ``**bold**``, ``*italic*``
  - native Word table for the Sumário GRI (read from ``Report.gri_index`` when
    present; otherwise falls back to the markdown already embedded in the
    ``sumario-gri`` section)

No branding, no diagramação, no PDF. Consultants open the file in Word and
apply the final layout.
"""

from __future__ import annotations

import io
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.models import Project, Report

_INLINE_GRI_PATTERN = re.compile(r"\(GRI\s+\d{1,3}-\d+[a-z]?\)", re.IGNORECASE)
_BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_PATTERN = re.compile(r"(?<!\*)\*([^*\n]+)\*(?!\*)")


def _family_order(rows: list[dict[str, Any]]) -> list[str]:
    families_seen: list[str] = []
    preferred = ["2", "3", "200", "300", "400"]
    grouped = {row.get("family") or "outros" for row in rows}
    for family in preferred:
        if family in grouped:
            families_seen.append(family)
    for family in sorted(grouped):
        if family not in preferred:
            families_seen.append(family)
    return families_seen


def _add_runs_with_inline_formatting(paragraph, text: str) -> None:
    """Emit text runs, turning ``**bold**`` into bold runs, ``*italic*`` into
    italic runs, and GRI parentheticals like ``(GRI 305-1)`` into italic runs.
    """
    # protect GRI parentheticals by wrapping them as italic markers
    text = _INLINE_GRI_PATTERN.sub(lambda m: f"*{m.group(0)}*", text)

    # tokenize by bold first, then by italic within non-bold chunks
    parts: list[tuple[str, str]] = []  # (style, text)
    last_end = 0
    for match in _BOLD_PATTERN.finditer(text):
        if match.start() > last_end:
            parts.append(("plain", text[last_end : match.start()]))
        parts.append(("bold", match.group(1)))
        last_end = match.end()
    if last_end < len(text):
        parts.append(("plain", text[last_end:]))

    for style, chunk in parts:
        if style == "bold":
            run = paragraph.add_run(chunk)
            run.bold = True
            continue
        # further split plain chunk on italic markers
        last = 0
        for it in _ITALIC_PATTERN.finditer(chunk):
            if it.start() > last:
                paragraph.add_run(chunk[last : it.start()])
            run = paragraph.add_run(it.group(1))
            run.italic = True
            last = it.end()
        if last < len(chunk):
            paragraph.add_run(chunk[last:])


def _is_markdown_table_line(line: str) -> bool:
    return "|" in line and line.strip().startswith("|")


def _parse_markdown_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    max_cols = 0
    for line in lines:
        stripped = line.strip().strip("|")
        if not stripped:
            continue
        if set(stripped.replace("|", "").strip()) <= {"-", ":", " "}:
            # separator row "---|---"
            continue
        cells = [cell.strip() for cell in stripped.split("|")]
        rows.append(cells)
        max_cols = max(max_cols, len(cells))
    # normalize: pad short rows with empty strings
    for row in rows:
        while len(row) < max_cols:
            row.append("")
    return rows


def _add_heading(document: Document, text: str, level: int) -> None:
    document.add_heading(text, level=level)


def _render_section_content(document: Document, content: str) -> None:
    """Emit the section content as a sequence of docx paragraphs/tables.

    Very small Markdown subset:
      - blank line separates blocks
      - ``### Title`` / ``## Title`` → sub-heading levels
      - ``- `` or ``* `` line → bullet paragraph
      - lines with leading ``|`` → markdown table
      - otherwise → normal paragraph with inline formatting
    """
    if not content:
        return

    lines = content.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        # markdown table block
        if _is_markdown_table_line(stripped):
            table_lines: list[str] = []
            while i < len(lines) and _is_markdown_table_line(lines[i].strip()):
                table_lines.append(lines[i])
                i += 1
            rows = _parse_markdown_table(table_lines)
            if rows:
                table = document.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Light Grid Accent 1"
                for r_idx, row_cells in enumerate(rows):
                    for c_idx, cell_text in enumerate(row_cells):
                        cell = table.rows[r_idx].cells[c_idx]
                        cell.text = ""
                        paragraph = cell.paragraphs[0]
                        _add_runs_with_inline_formatting(paragraph, cell_text)
            continue

        # sub-headings
        heading_match = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading_match:
            level = min(4, len(heading_match.group(1)))
            document.add_heading(heading_match.group(2), level=level)
            i += 1
            continue

        # bullet list block — gather consecutive bullet lines
        if re.match(r"^[-*]\s+", stripped):
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i].strip()):
                bullet_text = re.sub(r"^[-*]\s+", "", lines[i].strip())
                paragraph = document.add_paragraph(style="List Bullet")
                _add_runs_with_inline_formatting(paragraph, bullet_text)
                i += 1
            continue

        # gather a paragraph block (consecutive non-special non-empty lines)
        paragraph_lines: list[str] = []
        while i < len(lines):
            current = lines[i]
            stripped_current = current.strip()
            if not stripped_current:
                break
            if _is_markdown_table_line(stripped_current):
                break
            if re.match(r"^(#{2,4}|[-*])\s+", stripped_current):
                break
            paragraph_lines.append(stripped_current)
            i += 1
        if paragraph_lines:
            joined = " ".join(paragraph_lines)
            paragraph = document.add_paragraph()
            _add_runs_with_inline_formatting(paragraph, joined)


def _render_sumario_gri_table(
    document: Document, gri_index: list[dict[str, Any]]
) -> None:
    if not gri_index:
        return
    by_family: dict[str, list[dict[str, Any]]] = {}
    for row in gri_index:
        by_family.setdefault(row.get("family", "") or "outros", []).append(row)

    for family in _family_order(gri_index):
        family_rows = by_family.get(family, [])
        document.add_heading(f"GRI {family}", level=2)
        table = document.add_table(rows=1 + len(family_rows), cols=4)
        table.style = "Light Grid Accent 1"
        header_cells = table.rows[0].cells
        for idx, header in enumerate(
            ("Código", "Divulgação", "Evidência / Localização", "Status")
        ):
            header_cells[idx].paragraphs[0].add_run(header).bold = True
        for r_idx, row in enumerate(family_rows, start=1):
            cells = table.rows[r_idx].cells
            cells[0].text = row.get("code", "")
            cells[1].text = row.get("standard_text", "")
            excerpt = row.get("evidence_excerpt") or "—"
            section_ref = row.get("section_ref") or "—"
            cells[2].text = f"{excerpt} ({section_ref})"
            cells[3].text = row.get("status", "nao_atendido")


def generate_report_docx(report: Report, project: Project) -> bytes:
    document = Document()

    # base style tuning
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    # cover
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"Relatório de Sustentabilidade — {project.org_name}")
    title_run.bold = True
    title_run.font.size = Pt(20)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"Ano-base {project.base_year}")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(12)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"Versão preliminar v{report.version} — gerada automaticamente "
        "pelo Agente de IA ESG. Revisão humana obrigatória antes da publicação."
    )
    meta_run.font.size = Pt(10)

    document.add_page_break()

    sections_raw = report.sections
    sections = sections_raw if isinstance(sections_raw, list) else []
    ordered = sorted(
        [s for s in sections if isinstance(s, dict)],
        key=lambda s: s.get("order", 0),
    )
    gri_index_raw = report.gri_index
    gri_index = gri_index_raw if isinstance(gri_index_raw, list) else None

    for section in ordered:
        key = section.get("key")
        title = section.get("title", "")
        level = min(3, max(1, int(section.get("heading_level", 1) or 1)))
        document.add_heading(title, level=level)
        if key == "sumario-gri" and gri_index:
            # render the structured table rather than the embedded markdown
            _render_sumario_gri_table(document, gri_index)
        else:
            content = str(section.get("content", "") or "")
            _render_section_content(document, content)
        document.add_page_break()

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
