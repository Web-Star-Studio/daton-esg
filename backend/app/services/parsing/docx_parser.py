from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document as DocxDocument

from app.services.parsing import ParsedDocumentResult


def _normalize_text(value: str | None) -> str | None:
    if value is None:
        return None

    normalized = value.strip()
    return normalized or None


def parse_docx_document(file_bytes: bytes) -> ParsedDocumentResult:
    doc = DocxDocument(BytesIO(file_bytes))
    blocks: list[dict[str, Any]] = []
    extracted_sections: list[str] = []

    for paragraph in doc.paragraphs:
        text = _normalize_text(paragraph.text)
        if not text:
            continue

        style_name = (paragraph.style.name or "").lower()
        if style_name.startswith("heading"):
            heading_level = style_name.removeprefix("heading").strip() or "1"
            block_type = "heading"
            blocks.append(
                {
                    "type": block_type,
                    "level": heading_level,
                    "text": text,
                }
            )
            extracted_sections.append(f"{'#' * int(heading_level)} {text}")
            continue

        if style_name.startswith("list"):
            blocks.append({"type": "list_item", "text": text})
            extracted_sections.append(f"- {text}")
            continue

        blocks.append({"type": "paragraph", "text": text})
        extracted_sections.append(text)

    tables_payload = []
    for table in doc.tables:
        normalized_rows = []
        for row in table.rows:
            normalized_row = [_normalize_text(cell.text) for cell in row.cells]
            if any(cell not in {None, ""} for cell in normalized_row):
                normalized_rows.append(normalized_row)

        if not normalized_rows:
            continue

        header = normalized_rows[0]
        rows = normalized_rows[1:] if len(normalized_rows) > 1 else []
        tables_payload.append({"header": header, "rows": rows})

        extracted_sections.append(" | ".join(cell or "" for cell in header).strip())
        for row in rows:
            extracted_sections.append(" | ".join(cell or "" for cell in row).strip())

    return ParsedDocumentResult(
        extracted_text="\n\n".join(
            section for section in extracted_sections if section
        ).strip(),
        parsed_payload={
            "blocks": blocks,
            "tables": tables_payload,
        },
    )
