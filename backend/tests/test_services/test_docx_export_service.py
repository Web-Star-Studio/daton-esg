"""Unit tests for the DOCX exporter. Validates the bytes look like a valid
.docx archive and that headings/tables are present when expected.
"""

from __future__ import annotations

import io
import zipfile
from datetime import datetime, timezone
from uuid import uuid4

from docx import Document

from app.models import Project, Report
from app.models.enums import OrganizationSize, ProjectStatus, ReportStatus
from app.services.docx_export_service import generate_report_docx


def _make_project() -> Project:
    return Project(
        id=uuid4(),
        user_id=uuid4(),
        org_name="Cooperlíquidos",
        org_sector="Logística",
        org_size=OrganizationSize.MEDIUM,
        org_location="Porto Alegre",
        base_year=2025,
        scope="Operações nacionais",
        status=ProjectStatus.PRELIMINARY_REPORT,
        created_at=datetime.now(timezone.utc),
        updated_at=datetime.now(timezone.utc),
    )


def _make_report(project: Project) -> Report:
    return Report(
        id=uuid4(),
        project_id=project.id,
        version=1,
        status=ReportStatus.DRAFT,
        sections=[
            {
                "key": "a-empresa",
                "title": "A Empresa",
                "order": 1,
                "heading_level": 1,
                "content": (
                    "A Cooperlíquidos é uma cooperativa de transportes e "
                    "logística (GRI 2-1). A operação abrange múltiplos estados "
                    "brasileiros.\n\n"
                    "## Estrutura Organizacional\n\n"
                    "A governança é exercida pelo Conselho de Administração.\n\n"
                    "- Diretoria executiva\n"
                    "- Conselho fiscal\n\n"
                    "Enquadramento ESG e normativo\n"
                    "- Pilares ESG: G\n"
                    "- GRI aplicável: GRI 2-1\n"
                ),
                "gri_codes_used": ["GRI 2-1"],
                "word_count": 52,
                "status": "completed",
            },
            {
                "key": "gestao-ambiental",
                "title": "Gestão Ambiental",
                "order": 4,
                "heading_level": 1,
                "content": (
                    "O consumo mensal de água é monitorado (GRI 303-3).\n\n"
                    "| Mês | Consumo (m³) |\n"
                    "|---|---|\n"
                    "| Jan | 1.234 |\n"
                    "| Fev | 1.400 |\n\n"
                    "Enquadramento ESG e normativo\n"
                    "- Pilares ESG: E\n"
                ),
                "gri_codes_used": ["GRI 303-3"],
                "word_count": 30,
                "status": "completed",
            },
            {
                "key": "sumario-gri",
                "title": "Sumário GRI",
                "order": 14,
                "heading_level": 1,
                "content": "",
                "gri_codes_used": [],
                "word_count": 0,
                "status": "completed",
            },
        ],
        gri_index=[
            {
                "code": "GRI 2-1",
                "family": "2",
                "standard_text": "Detalhes da organização",
                "evidence_excerpt": "cooperativa de transportes",
                "section_ref": "a-empresa",
                "status": "atendido",
                "found_in_text": True,
            },
            {
                "code": "GRI 303-3",
                "family": "300",
                "standard_text": "Captação de água",
                "evidence_excerpt": "consumo mensal monitorado",
                "section_ref": "gestao-ambiental",
                "status": "atendido",
                "found_in_text": True,
            },
            {
                "code": "GRI 2-22",
                "family": "2",
                "standard_text": "Estratégia de desenvolvimento",
                "evidence_excerpt": None,
                "section_ref": None,
                "status": "nao_atendido",
                "found_in_text": False,
            },
        ],
        gaps=[],
        created_at=datetime.now(timezone.utc),
        updated_at=datetime.now(timezone.utc),
    )


def test_generate_report_docx_returns_valid_docx_bytes() -> None:
    project = _make_project()
    report = _make_report(project)

    payload = generate_report_docx(report, project)

    assert isinstance(payload, bytes)
    assert payload[:2] == b"PK"  # DOCX is a ZIP archive
    assert zipfile.is_zipfile(io.BytesIO(payload))


def test_generated_docx_has_expected_headings_and_tables() -> None:
    project = _make_project()
    report = _make_report(project)
    payload = generate_report_docx(report, project)

    document = Document(io.BytesIO(payload))
    heading_texts = [
        p.text
        for p in document.paragraphs
        if p.style and p.style.name.startswith("Heading")
    ]
    assert "A Empresa" in heading_texts
    assert "Gestão Ambiental" in heading_texts
    assert "Sumário GRI" in heading_texts

    # Sumário GRI should render as tables grouped by family, not as plain text
    assert len(document.tables) >= 2

    all_text = "\n".join(p.text for p in document.paragraphs)
    assert "Cooperlíquidos" in all_text


def test_generate_docx_handles_empty_sections() -> None:
    project = _make_project()
    report = _make_report(project)
    report.sections = []
    report.gri_index = None

    payload = generate_report_docx(report, project)
    assert payload[:2] == b"PK"
    document = Document(io.BytesIO(payload))
    all_text = "\n".join(p.text for p in document.paragraphs)
    assert "Cooperlíquidos" in all_text  # cover still rendered
