from __future__ import annotations

from contextlib import AbstractAsyncContextManager
from io import BytesIO
from types import SimpleNamespace
from uuid import uuid4

import pytest
from docx import Document as DocxDocument
from openpyxl import Workbook
from sqlalchemy.dialects import postgresql

from app.models import Document, Project
from app.models.enums import (
    DocumentFileType,
    DocumentParsingStatus,
    ProjectStatus,
    UserRole,
)
from app.models.user import User
from app.services.parsing import ParsedDocumentResult
from app.services.parsing.docx_parser import parse_docx_document
from app.services.parsing.excel_parser import (
    parse_csv_document,
    parse_xlsx_document,
)
from app.services.parsing.orchestrator import run_document_parsing
from app.services.parsing.pdf_parser import (
    _build_textract_client,
    parse_pdf_document,
)
from app.services.storage_service import StorageService


class FakeStorage:
    def __init__(self, payload: bytes) -> None:
        self.payload = payload
        self.keys: list[str] = []
        self.bucket_name = "documents-bucket"
        self.settings = SimpleNamespace(
            document_parsing_pdf_provider="textract",
            environment="production",
            aws_endpoint_url="",
        )

    async def get_object_bytes(self, *, key: str) -> bytes:
        self.keys.append(key)
        return self.payload


class ExplodingStorage(FakeStorage):
    async def get_object_bytes(self, *, key: str) -> bytes:
        raise AssertionError(f"get_object_bytes should not be called for {key}")


class FakeExecuteResult:
    def __init__(self, document: Document | None) -> None:
        self.document = document

    def scalar_one_or_none(self) -> Document | None:
        return self.document


class FakeSession:
    def __init__(self, document: Document | None) -> None:
        self.document = document
        self.commit_count = 0
        self.rollback_count = 0
        self.statements = []

    async def execute(self, _statement):
        self.statements.append(_statement)
        return FakeExecuteResult(self.document)

    async def commit(self) -> None:
        self.commit_count += 1

    async def rollback(self) -> None:
        self.rollback_count += 1

    def begin(self):
        return FakeTransaction(self)


class FakeTransaction(AbstractAsyncContextManager[None]):
    def __init__(self, session: FakeSession) -> None:
        self.session = session

    async def __aenter__(self) -> None:
        return None

    async def __aexit__(self, exc_type, exc, tb) -> None:
        if exc_type is None:
            self.session.commit_count += 1
        return None


class FakeSessionContext(AbstractAsyncContextManager[FakeSession]):
    def __init__(self, session: FakeSession) -> None:
        self.session = session

    async def __aenter__(self) -> FakeSession:
        return self.session

    async def __aexit__(self, exc_type, exc, tb) -> None:
        return None


def make_project() -> Project:
    user = User(
        id=uuid4(),
        cognito_sub="cognito-sub-1",
        email="consultor@example.com",
        name="Consultor ESG",
        role=UserRole.CONSULTANT,
    )
    return Project(
        id=uuid4(),
        user_id=user.id,
        org_name="Acme Inc.",
        base_year=2025,
        status=ProjectStatus.COLLECTING,
    )


def make_document(project: Project, *, file_type: DocumentFileType) -> Document:
    document_id = uuid4()
    extension = file_type.value
    return Document(
        id=document_id,
        project_id=project.id,
        filename=f"evidencia.{extension}",
        file_type=file_type,
        s3_key=f"uploads/{project.id}/{document_id}/evidencia.{extension}",
        file_size_bytes=1024,
        parsing_status=DocumentParsingStatus.PENDING,
    )


def test_parse_csv_document_returns_structured_payload() -> None:
    result = parse_csv_document(
        b"Indicador,Valor,Unidade\nEnergia,1500,kWh\nAgua,200,m3\n"
    )

    assert "Energia" in result.extracted_text
    assert result.parsed_payload["format"] == "csv"
    assert result.parsed_payload["header"] == ["Indicador", "Valor", "Unidade"]
    assert result.parsed_payload["rows"][0] == ["Energia", "1500", "kWh"]


def test_parse_xlsx_document_reads_all_sheets() -> None:
    workbook = Workbook()
    active_sheet = workbook.active
    active_sheet.title = "Energia"
    active_sheet.append(["Indicador", "Valor"])
    active_sheet.append(["Consumo", 1500])
    second_sheet = workbook.create_sheet("Agua")
    second_sheet.append(["Indicador", "Valor"])
    second_sheet.append(["Consumo", 200])

    buffer = BytesIO()
    workbook.save(buffer)

    result = parse_xlsx_document(buffer.getvalue())

    assert result.parsed_payload["format"] == "xlsx"
    assert len(result.parsed_payload["sheets"]) == 2
    assert result.parsed_payload["sheets"][0]["sheet_name"] == "Energia"
    assert "Aba: Agua" in result.extracted_text


def test_parse_docx_document_preserves_structure() -> None:
    document = DocxDocument()
    document.add_heading("Política Ambiental", level=1)
    document.add_paragraph("A companhia mantém metas anuais.")
    document.add_paragraph("Reduzir emissões", style="List Bullet")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Indicador"
    table.rows[0].cells[1].text = "Valor"
    table.rows[1].cells[0].text = "Energia"
    table.rows[1].cells[1].text = "1500"

    buffer = BytesIO()
    document.save(buffer)

    result = parse_docx_document(buffer.getvalue())

    assert "Política Ambiental" in result.extracted_text
    assert result.parsed_payload["blocks"][0]["type"] == "heading"
    assert result.parsed_payload["tables"][0]["header"] == ["Indicador", "Valor"]


def test_parse_docx_document_defaults_heading_level_for_non_numeric_suffix() -> None:
    document = DocxDocument()
    paragraph = document.add_paragraph("Politica de Fornecedores")
    paragraph.style = document.styles["Heading 1"]
    paragraph.style.name = "Heading Annex"

    buffer = BytesIO()
    document.save(buffer)

    result = parse_docx_document(buffer.getvalue())

    assert result.parsed_payload["blocks"][0]["level"] == "1"
    assert result.extracted_text.startswith("# Politica de Fornecedores")


@pytest.mark.asyncio
async def test_parse_pdf_document_uses_local_provider_when_configured(
    monkeypatch,
) -> None:
    expected_result = ParsedDocumentResult(
        extracted_text="texto do pdf",
        parsed_payload={"provider": "local", "tables": []},
    )

    def fake_parse_pdf_locally(_file_bytes: bytes) -> ParsedDocumentResult:
        return expected_result

    monkeypatch.setattr(
        "app.services.parsing.pdf_parser.parse_pdf_locally",
        fake_parse_pdf_locally,
    )

    result = await parse_pdf_document(
        file_bytes=b"%PDF-1.4",
        settings=SimpleNamespace(document_parsing_pdf_provider="local"),
    )

    assert result == expected_result


def test_should_use_local_pdf_parser_handles_localhost_endpoints() -> None:
    from app.services.parsing.pdf_parser import should_use_local_pdf_parser

    settings = SimpleNamespace(
        document_parsing_pdf_provider="auto",
        aws_endpoint_url="http://localhost:4566",
        environment="production",
    )

    assert should_use_local_pdf_parser(settings) is True


def test_should_use_local_pdf_parser_handles_loopback_endpoints() -> None:
    from app.services.parsing.pdf_parser import should_use_local_pdf_parser

    settings = SimpleNamespace(
        document_parsing_pdf_provider="auto",
        aws_endpoint_url="http://127.0.0.1:4566",
        environment="production",
    )

    assert should_use_local_pdf_parser(settings) is True


def test_build_textract_client_uses_dedicated_textract_region(monkeypatch) -> None:
    captured_kwargs = {}

    def fake_boto3_client(**kwargs):
        captured_kwargs.update(kwargs)
        return object()

    monkeypatch.setattr(
        "app.services.parsing.pdf_parser.boto3.client",
        fake_boto3_client,
    )

    _build_textract_client(
        SimpleNamespace(
            aws_region="sa-east-1",
            aws_textract_region="us-east-1",
            aws_access_key_id="test",
            aws_secret_access_key="test",
            aws_endpoint_url="",
        )
    )

    assert captured_kwargs["service_name"] == "textract"
    assert captured_kwargs["region_name"] == "us-east-1"


@pytest.mark.asyncio
async def test_parse_pdf_document_uses_async_textract_and_aggregates_pages(
    monkeypatch,
) -> None:
    class FakeTextractClient:
        def __init__(self) -> None:
            self.start_calls = []
            self.analysis_calls = []

        def start_document_analysis(self, **kwargs):
            self.start_calls.append(kwargs)
            return {"JobId": "job-123"}

        def get_document_analysis(self, **kwargs):
            self.analysis_calls.append(kwargs)
            if len(self.analysis_calls) == 1:
                return {"JobStatus": "IN_PROGRESS"}
            if len(self.analysis_calls) == 2:
                return {
                    "JobStatus": "SUCCEEDED",
                    "Blocks": [
                        {
                            "Id": "line-1",
                            "BlockType": "LINE",
                            "Page": 1,
                            "Text": "Politica Ambiental",
                            "Geometry": {"BoundingBox": {"Top": 0.1}},
                        },
                        {
                            "Id": "table-1",
                            "BlockType": "TABLE",
                            "Page": 1,
                            "Relationships": [{"Type": "CHILD", "Ids": ["cell-1"]}],
                        },
                        {
                            "Id": "cell-1",
                            "BlockType": "CELL",
                            "RowIndex": 1,
                            "ColumnIndex": 1,
                            "Relationships": [{"Type": "CHILD", "Ids": ["word-1"]}],
                        },
                        {"Id": "word-1", "BlockType": "WORD", "Text": "Indicador"},
                    ],
                    "NextToken": "next-1",
                }
            return {
                "JobStatus": "SUCCEEDED",
                "Blocks": [
                    {
                        "Id": "line-2",
                        "BlockType": "LINE",
                        "Page": 2,
                        "Text": "Consumo de Energia",
                        "Geometry": {"BoundingBox": {"Top": 0.2}},
                    }
                ],
            }

    client = FakeTextractClient()
    sleep_calls: list[float] = []

    async def fake_sleep(seconds: float) -> None:
        sleep_calls.append(seconds)

    monkeypatch.setattr(
        "app.services.parsing.pdf_parser._build_textract_client",
        lambda _settings: client,
    )
    monkeypatch.setattr("app.services.parsing.pdf_parser.asyncio.sleep", fake_sleep)

    result = await parse_pdf_document(
        bucket_name="documents-bucket",
        key="uploads/project/doc.pdf",
        settings=SimpleNamespace(
            document_parsing_pdf_provider="textract",
            aws_region="sa-east-1",
            aws_textract_region="us-east-1",
        ),
    )

    assert client.start_calls == [
        {
            "DocumentLocation": {
                "S3Object": {
                    "Bucket": "documents-bucket",
                    "Name": "uploads/project/doc.pdf",
                }
            },
            "FeatureTypes": ["TABLES"],
        }
    ]
    assert sleep_calls == [2.0]
    assert result.parsed_payload["provider"] == "textract"
    assert result.parsed_payload["job_id"] == "job-123"
    assert len(result.parsed_payload["pages"]) == 2
    assert result.parsed_payload["tables"][0]["header"] == ["Indicador"]
    assert "Politica Ambiental" in result.extracted_text
    assert "Consumo de Energia" in result.extracted_text


@pytest.mark.asyncio
async def test_parse_pdf_document_raises_when_textract_job_fails(monkeypatch) -> None:
    class FakeTextractClient:
        def start_document_analysis(self, **_kwargs):
            return {"JobId": "job-456"}

        def get_document_analysis(self, **_kwargs):
            return {"JobStatus": "FAILED", "StatusMessage": "Unsupported document"}

    monkeypatch.setattr(
        "app.services.parsing.pdf_parser._build_textract_client",
        lambda _settings: FakeTextractClient(),
    )

    with pytest.raises(RuntimeError, match="Unsupported document"):
        await parse_pdf_document(
            bucket_name="documents-bucket",
            key="uploads/project/doc.pdf",
            settings=SimpleNamespace(
                document_parsing_pdf_provider="textract",
                aws_region="sa-east-1",
                aws_textract_region="us-east-1",
            ),
        )


@pytest.mark.asyncio
async def test_run_document_parsing_marks_document_completed(monkeypatch) -> None:
    project = make_project()
    document = make_document(project, file_type=DocumentFileType.CSV)
    session = FakeSession(document)
    storage = FakeStorage(b"Indicador,Valor\nEnergia,1500\n")

    monkeypatch.setattr(
        "app.services.parsing.orchestrator.SessionLocal",
        lambda: FakeSessionContext(session),
    )

    await run_document_parsing(document.id, storage=storage)

    assert storage.keys == [document.s3_key]
    assert document.parsing_status == DocumentParsingStatus.COMPLETED
    assert document.parsed_payload is not None
    assert "Energia" in (document.extracted_text or "")
    assert session.commit_count >= 2
    compiled = session.statements[0].compile(dialect=postgresql.dialect())
    assert "FOR UPDATE" in str(compiled)


@pytest.mark.asyncio
async def test_run_document_parsing_uses_s3_textract_path_for_pdf(monkeypatch) -> None:
    project = make_project()
    document = make_document(project, file_type=DocumentFileType.PDF)
    session = FakeSession(document)
    storage = ExplodingStorage(b"%PDF-1.4")
    called_kwargs: dict[str, object] = {}

    async def fake_parse_document(**kwargs):
        called_kwargs.update(kwargs)
        return ParsedDocumentResult(
            extracted_text="Texto OCR",
            parsed_payload={"provider": "textract", "job_id": "job-789", "tables": []},
        )

    monkeypatch.setattr(
        "app.services.parsing.orchestrator.SessionLocal",
        lambda: FakeSessionContext(session),
    )
    monkeypatch.setattr(
        "app.services.parsing.orchestrator.parse_pdf_document",
        fake_parse_document,
    )

    await run_document_parsing(document.id, storage=storage)

    assert called_kwargs["bucket_name"] == storage.bucket_name
    assert called_kwargs["key"] == document.s3_key
    assert called_kwargs["settings"] is storage.settings
    assert document.parsing_status == DocumentParsingStatus.COMPLETED
    assert document.extracted_text == "Texto OCR"
    assert document.parsed_payload["job_id"] == "job-789"


@pytest.mark.asyncio
async def test_run_document_parsing_uses_local_pdf_bytes_when_configured(
    monkeypatch,
) -> None:
    project = make_project()
    document = make_document(project, file_type=DocumentFileType.PDF)
    session = FakeSession(document)
    storage = FakeStorage(b"%PDF-1.4")

    async def fake_parse_document(**kwargs):
        assert kwargs["file_bytes"] == b"%PDF-1.4"
        assert kwargs["settings"] is storage.settings
        return ParsedDocumentResult(
            extracted_text="Texto local",
            parsed_payload={"provider": "local", "tables": []},
        )

    storage.settings = SimpleNamespace(document_parsing_pdf_provider="local")

    monkeypatch.setattr(
        "app.services.parsing.orchestrator.SessionLocal",
        lambda: FakeSessionContext(session),
    )
    monkeypatch.setattr(
        "app.services.parsing.orchestrator.parse_pdf_document",
        fake_parse_document,
    )

    await run_document_parsing(document.id, storage=storage)

    assert storage.keys == [document.s3_key]
    assert document.parsing_status == DocumentParsingStatus.COMPLETED
    assert document.extracted_text == "Texto local"


@pytest.mark.asyncio
async def test_run_document_parsing_returns_when_document_is_already_processing(
    monkeypatch,
) -> None:
    project = make_project()
    document = make_document(project, file_type=DocumentFileType.CSV)
    document.parsing_status = DocumentParsingStatus.PROCESSING
    session = FakeSession(document)
    storage = FakeStorage(b"Indicador,Valor\nEnergia,1500\n")
    parse_called = False

    async def fake_parse_document(_document: Document, *, storage_service):
        nonlocal parse_called
        parse_called = True
        return ParsedDocumentResult(extracted_text="", parsed_payload={})

    monkeypatch.setattr(
        "app.services.parsing.orchestrator.SessionLocal",
        lambda: FakeSessionContext(session),
    )
    monkeypatch.setattr(
        "app.services.parsing.orchestrator._parse_document_by_type",
        fake_parse_document,
    )

    await run_document_parsing(document.id, storage=storage)

    assert parse_called is False
    assert session.commit_count == 1
    compiled = session.statements[0].compile(dialect=postgresql.dialect())
    assert "FOR UPDATE" in str(compiled)


@pytest.mark.asyncio
async def test_run_document_parsing_returns_when_document_is_already_completed(
    monkeypatch,
) -> None:
    project = make_project()
    document = make_document(project, file_type=DocumentFileType.CSV)
    document.parsing_status = DocumentParsingStatus.COMPLETED
    session = FakeSession(document)
    storage = FakeStorage(b"Indicador,Valor\nEnergia,1500\n")
    parse_called = False

    async def fake_parse_document(_document: Document, *, storage_service):
        nonlocal parse_called
        parse_called = True
        return ParsedDocumentResult(extracted_text="", parsed_payload={})

    monkeypatch.setattr(
        "app.services.parsing.orchestrator.SessionLocal",
        lambda: FakeSessionContext(session),
    )
    monkeypatch.setattr(
        "app.services.parsing.orchestrator._parse_document_by_type",
        fake_parse_document,
    )

    await run_document_parsing(document.id, storage=storage)

    assert parse_called is False
    assert session.commit_count == 1
    compiled = session.statements[0].compile(dialect=postgresql.dialect())
    assert "FOR UPDATE" in str(compiled)


@pytest.mark.asyncio
async def test_run_document_parsing_retries_once_then_fails(monkeypatch) -> None:
    project = make_project()
    document = make_document(project, file_type=DocumentFileType.CSV)
    document.extracted_text = "Texto anterior"
    document.parsed_payload = {"provider": "local"}
    session = FakeSession(document)
    storage = FakeStorage(b"Indicador,Valor\nEnergia,1500\n")
    attempts = {"count": 0}

    async def fake_parse_document(_document: Document, *, storage_service):
        assert storage_service is storage
        attempts["count"] += 1
        raise RuntimeError(f"falha {attempts['count']}")

    monkeypatch.setattr(
        "app.services.parsing.orchestrator.SessionLocal",
        lambda: FakeSessionContext(session),
    )
    monkeypatch.setattr(
        "app.services.parsing.orchestrator._parse_document_by_type",
        fake_parse_document,
    )

    await run_document_parsing(document.id, storage=storage)

    assert attempts["count"] == 2
    assert session.rollback_count == 2
    assert document.parsing_status == DocumentParsingStatus.FAILED
    assert document.extracted_text is None
    assert document.parsed_payload is None
    assert document.parsing_error == "Parsing failed"


def test_storage_service_get_object_bytes_closes_streaming_body() -> None:
    class FakeBody:
        def __init__(self) -> None:
            self.closed = False

        def read(self) -> bytes:
            return b"payload"

        def close(self) -> None:
            self.closed = True

    class FakeClient:
        def __init__(self, body: FakeBody) -> None:
            self.body = body

        def get_object(self, **_kwargs):
            return {"Body": self.body}

    body = FakeBody()
    service = StorageService.__new__(StorageService)
    service.bucket_name = "documents-bucket"
    service._client = FakeClient(body)

    result = service._get_object_bytes(key="uploads/project/doc.pdf")

    assert result == b"payload"
    assert body.closed is True
